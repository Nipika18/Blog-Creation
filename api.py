from fastapi import FastAPI, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from sqlalchemy.exc import OperationalError
from sqlalchemy import text
from datetime import timedelta, date
import models, database, auth
from pydantic import BaseModel
from typing import Optional, List
import os
import re
from pathlib import Path
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse, HTMLResponse
import io
from fastapi.security import OAuth2PasswordBearer
import secrets
import requests
import json
from datetime import datetime, timedelta

# Import the blog writer backend
from bwa_backend import app as blog_app


app = FastAPI(title="Blog Writing Agent API")

# Enable CORS for React frontend
origins = [
    "http://localhost:5173",
    "http://127.0.0.1:5173",
    "http://localhost:3000",
    "http://localhost:8000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"],
)


# Create tables
models.Base.metadata.create_all(bind=database.engine)


class UserCreate(BaseModel):
    email: str
    password: str

class Token(BaseModel):
    access_token: str
    token_type: str

class ForgotPasswordRequest(BaseModel):
    email: str

class PasswordResetConfirm(BaseModel):
    email: str
    new_password: str

# Security
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="login")

def get_current_user(db: Session = Depends(database.get_db), token: str = Depends(oauth2_scheme)):
    try:
        payload = auth.jwt.decode(token, auth.SECRET_KEY, algorithms=[auth.ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise HTTPException(status_code=401, detail="Invalid token")
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")
        
    user = db.query(models.User).filter(models.User.email == email).first()
    if user is None:
        raise HTTPException(status_code=401, detail="User not found")
    return user

# --- Image Management ---

@app.get("/images/{filename}")
def get_image(filename: str, db: Session = Depends(database.get_db)):
    db_image = db.query(models.BlogImage).filter(models.BlogImage.filename == filename).first()
    if not db_image:
        raise HTTPException(status_code=404, detail="Image not found")
    
    ext = os.path.splitext(filename)[1].lower()
    media_type = "image/png"
    if ext == ".jpg" or ext == ".jpeg": media_type = "image/jpeg"
    elif ext == ".webp": media_type = "image/webp"
    elif ext == ".gif": media_type = "image/gif"
        
    return StreamingResponse(io.BytesIO(db_image.content), media_type=media_type)


# --- Public Blog Page (for LinkedIn Article sharing) ---

@app.get("/p/{filename}", response_class=HTMLResponse)
def public_blog_page(filename: str, db: Session = Depends(database.get_db)):
    """Serves a beautiful public HTML page for a blog — used as the LinkedIn article link."""
    blog = db.query(models.Blog).filter(models.Blog.filename == filename).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    
    import markdown
    blog_html = markdown.markdown(blog.content, extensions=['extra', 'codehilite', 'toc'])
    
    # Fix image paths to absolute URLs
    blog_html = re.sub(r'src="/?images/', f'src="/images/', blog_html)
    
    # Get public URL (default to live Render domain)
    public_base_url = os.getenv("PUBLIC_URL", "https://blog-creation-eetn.onrender.com").rstrip('/')
    
    # Extract first image for OG meta tag
    og_image = ""
    img_match = re.search(r'src=["\']([^"\']+)["\']', blog_html)
    if img_match:
        img_src = img_match.group(1)
        if img_src.startswith('http://') or img_src.startswith('https://'):
            og_image = img_src
        else:
            if not img_src.startswith('/'):
                img_src = f"/{img_src}"
            og_image = f"{public_base_url}{img_src}"
    
    # Extract description (first 200 chars of text)
    text_only = re.sub(r'<[^>]+>', '', blog_html)
    description = text_only[:200].strip().replace('"', '&quot;').replace('\n', ' ')
    
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{blog.title} | Blog Studio</title>
    <meta property="og:title" content="{blog.title}" />
    <meta property="og:description" content="{description}" />
    <meta property="og:type" content="article" />
    <meta property="og:url" content="{public_base_url}/p/{filename}" />
    {f'<meta property="og:image" content="{og_image}" />' if og_image else ''}
    {f'<meta property="og:image:secure_url" content="{og_image}" />' if og_image else ''}
    {f'<meta name="twitter:image" content="{og_image}" />' if og_image else ''}
    <meta name="twitter:card" content="summary_large_image">
    <meta name="author" content="Blog Studio">
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=Merriweather:wght@400;700&display=swap" rel="stylesheet">
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Merriweather', Georgia, serif;
            background: #f4f6f8;
            color: #1a1a2e;
            line-height: 1.85;
            padding: 2rem 1rem;
        }}
        .hero-banner {{
            width: 100%;
            max-height: 420px;
            object-fit: cover;
            display: block;
        }}
        .article-container {{
            max-width: 880px;
            margin: 0 auto;
            padding: 3.5rem 3rem 4.5rem;
            background: #ffffff;
            border-radius: 16px;
            box-shadow: 0 4px 24px rgba(0, 0, 0, 0.06);
            border: 1px solid #e9ecef;
        }}
        h1 {{
            font-family: 'Inter', sans-serif;
            font-size: 2.4rem;
            font-weight: 700;
            line-height: 1.25;
            margin-bottom: 0.5rem;
            color: #0a0a23;
        }}
        .article-meta {{
            font-family: 'Inter', sans-serif;
            font-size: 0.85rem;
            color: #6b7280;
            margin-bottom: 2rem;
            padding-bottom: 1.5rem;
            border-bottom: 1px solid #e5e7eb;
        }}
        h2 {{
            font-family: 'Inter', sans-serif;
            font-size: 1.5rem;
            font-weight: 600;
            margin-top: 2.5rem;
            margin-bottom: 1rem;
            color: #1e3a5f;
        }}
        h3 {{
            font-family: 'Inter', sans-serif;
            font-size: 1.2rem;
            font-weight: 600;
            margin-top: 2rem;
            margin-bottom: 0.75rem;
        }}
        p {{
            font-size: 1.05rem;
            margin-bottom: 1.4rem;
        }}
        ul, ol {{
            margin: 1rem 0 1.5rem 1.5rem;
        }}
        li {{
            margin-bottom: 0.6rem;
            font-size: 1rem;
        }}
        img {{
            max-width: 100%;
            height: auto;
            border-radius: 12px;
            margin: 2rem 0;
            box-shadow: 0 4px 20px rgba(0,0,0,0.08);
        }}
        blockquote {{
            border-left: 4px solid #3b82f6;
            padding: 1rem 1.5rem;
            margin: 1.5rem 0;
            background: #eff6ff;
            border-radius: 0 8px 8px 0;
            font-style: italic;
            color: #374151;
        }}
        code {{
            background: #f1f5f9;
            padding: 2px 6px;
            border-radius: 4px;
            font-size: 0.9rem;
        }}
        pre {{
            background: #1e293b;
            color: #e2e8f0;
            padding: 1.25rem;
            border-radius: 10px;
            overflow-x: auto;
            margin: 1.5rem 0;
        }}
        pre code {{
            background: none;
            padding: 0;
            color: inherit;
        }}
        @media (max-width: 768px) {{
            body {{
                padding: 0.75rem 0.5rem;
                background: #f8f9fa;
            }}
            .article-container {{
                padding: 2rem 1.25rem 3rem;
                border-radius: 12px;
                width: 100%;
            }}
            h1 {{
                font-size: 1.85rem;
            }}
            h2 {{
                font-size: 1.35rem;
                margin-top: 2rem;
            }}
            h3 {{
                font-size: 1.15rem;
            }}
            p, li {{
                font-size: 1rem;
                line-height: 1.75;
            }}
            img {{
                margin: 1.25rem 0;
                border-radius: 8px;
            }}
            blockquote {{
                padding: 0.75rem 1rem;
                margin: 1.25rem 0;
            }}
            pre {{
                padding: 1rem;
                font-size: 0.85rem;
            }}
        }}
        @media (max-width: 480px) {{
            body {{
                padding: 0;
                background: #ffffff;
            }}
            .article-container {{
                padding: 1.5rem 1rem 3rem;
                border-radius: 0;
                border: none;
                box-shadow: none;
            }}
            h1 {{
                font-size: 1.6rem;
            }}
        }}
        .footer {{
            text-align: center;
            padding: 2rem;
            font-family: 'Inter', sans-serif;
            font-size: 0.8rem;
            color: #9ca3af;
            border-top: 1px solid #e5e7eb;
            margin-top: 3rem;
        }}
        .footer a {{
            color: #3b82f6;
            text-decoration: none;
        }}
        em {{
            font-size: 0.85rem;
            color: #6b7280;
            display: block;
            text-align: center;
            margin-top: -0.75rem;
            margin-bottom: 2rem;
        }}
    </style>
</head>
<body>
    <div class="article-container">
        <h1>{blog.title}</h1>
        <div class="article-meta">Published on Blog Studio</div>
        {blog_html}
    </div>
    <div class="footer">
        Written with <a href="#">Blog Studio</a> &mdash; AI-Powered Blog Writing Agent
    </div>
</body>
</html>"""
    return HTMLResponse(content=html)

class BlogImageGenerateRequest(BaseModel):
    placeholder: str
    prompt: Optional[str] = None

@app.post("/blog/{filename}/generate-image")
async def generate_blog_image(
    filename: str,
    req: BlogImageGenerateRequest,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(get_current_user)
):
    blog = db.query(models.Blog).filter(
        models.Blog.filename == filename,
        models.Blog.user_id == current_user.id
    ).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")

    prompt = req.prompt
    placeholder = req.placeholder

    if not prompt and ":" in placeholder:
        parts = placeholder.split(":", 1)
        prompt = parts[1].rstrip("]").strip()

    if not prompt:
        prompt = blog.title

    from bwa_backend import generate_single_image_bytes
    try:
        img_bytes = generate_single_image_bytes(prompt, topic=blog.title)
    except Exception as e:
        print(f"Single image generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Image generation failed: {str(e)}")

    if not img_bytes:
        raise HTTPException(status_code=500, detail="Could not generate image for this prompt.")

    image_filename = f"{secrets.token_hex(8)}_ai.webp"
    db_image = models.BlogImage(
        filename=image_filename,
        content=img_bytes,
        blog_id=blog.id
    )
    db.add(db_image)

    # Replace placeholder tag in content with image markdown
    alt_text = prompt[:50]
    img_md = f"\n\n![{alt_text}](images/{image_filename})\n\n"

    content = blog.content
    num_match = re.search(r"\d+", placeholder)
    if num_match:
        num = num_match.group(0)
        pattern = rf"\[+IMAGE_?(?:PLACEHOLDER_)?{num}(?:\:[^\]]+)?\]+"
        content = re.sub(pattern, img_md, content, count=1)
    elif placeholder in content:
        content = content.replace(placeholder, img_md)

    blog.content = content
    db.commit()

    return {
        "success": True,
        "url": f"/images/{image_filename}",
        "filename": image_filename,
        "updated_content": blog.content
    }


class BlogImageActionRequest(BaseModel):
    image_filename: str


@app.post("/blog/{filename}/remove-image")
async def remove_blog_image(
    filename: str,
    req: BlogImageActionRequest,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(get_current_user)
):
    """Remove a generated image and restore the placeholder so user can regenerate or upload."""
    blog = db.query(models.Blog).filter(
        models.Blog.filename == filename,
        models.Blog.user_id == current_user.id
    ).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")

    # Find the image markdown pattern: ![alt text](images/xxx.webp) or ![alt](images/xxx.webp)
    img_pattern = rf'\n*!\[[^\]]*\]\([^)]*{re.escape(req.image_filename)}[^)]*\)\n*'
    match = re.search(img_pattern, blog.content)
    if not match:
        raise HTTPException(status_code=404, detail="Image not found in blog content")

    # Determine which image slot this was (count image markdowns before this one)
    preceding_content = blog.content[:match.start()]
    existing_images = re.findall(r'!\[[^\]]*\]\([^)]+\)', preceding_content)
    slot_num = len(existing_images) + 1

    # Restore the placeholder tag
    placeholder_tag = f"[[IMAGE_{slot_num}: {blog.title}]]"
    blog.content = blog.content[:match.start()] + f"\n\n{placeholder_tag}\n\n" + blog.content[match.end():]
    db.commit()

    return {
        "success": True,
        "updated_content": blog.content
    }


@app.post("/blog/{filename}/regenerate-image")
async def regenerate_blog_image(
    filename: str,
    req: BlogImageActionRequest,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(get_current_user)
):
    """Replace an existing generated image with a newly generated one."""
    blog = db.query(models.Blog).filter(
        models.Blog.filename == filename,
        models.Blog.user_id == current_user.id
    ).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")

    # Find the image markdown
    img_pattern = rf'!\[[^\]]*\]\([^)]*{re.escape(req.image_filename)}[^)]*\)'
    match = re.search(img_pattern, blog.content)
    if not match:
        raise HTTPException(status_code=404, detail="Image not found in blog content")

    # Extract the alt text as prompt
    alt_match = re.search(r'!\[([^\]]*)\]', match.group(0))
    prompt = alt_match.group(1) if alt_match else blog.title

    from bwa_backend import generate_single_image_bytes
    try:
        img_bytes = generate_single_image_bytes(prompt, topic=blog.title)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Image regeneration failed: {str(e)}")

    if not img_bytes:
        raise HTTPException(status_code=500, detail="Could not regenerate image.")

    new_image_filename = f"{secrets.token_hex(8)}_ai.webp"
    db_image = models.BlogImage(
        filename=new_image_filename,
        content=img_bytes,
        blog_id=blog.id
    )
    db.add(db_image)

    # Replace old image markdown with new one
    new_img_md = f"![{prompt[:50]}](images/{new_image_filename})"
    blog.content = blog.content[:match.start()] + new_img_md + blog.content[match.end():]
    db.commit()

    return {
        "success": True,
        "url": f"/images/{new_image_filename}",
        "filename": new_image_filename,
        "updated_content": blog.content
    }

@app.post("/blog/{filename}/upload-image")
async def upload_blog_image(
    filename: str, 
    image: UploadFile = File(...), 
    placeholder: Optional[str] = Form(None),
    db: Session = Depends(database.get_db), 
    current_user: models.User = Depends(get_current_user)
):
    blog = db.query(models.Blog).filter(models.Blog.filename == filename, models.Blog.user_id == current_user.id).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    
    content = await image.read()
    try:
        from bwa_backend import _resize_image_bytes
        content = _resize_image_bytes(content, "1024x576", quality=85)
    except Exception:
        pass

    ext = os.path.splitext(image.filename)[1].lower() or ".webp"
    image_filename = f"{secrets.token_hex(8)}{ext}"
    
    db_image = models.BlogImage(
        filename=image_filename,
        content=content,
        blog_id=blog.id
    )
    db.add(db_image)

    # Replace placeholder if provided
    img_md = f"\n\n![Uploaded Image](images/{image_filename})\n\n"
    if placeholder:
        num_match = re.search(r"\d+", placeholder)
        if num_match:
            num = num_match.group(0)
            pattern = rf"\[+IMAGE_?(?:PLACEHOLDER_)?{num}(?:\:[^\]]+)?\]+"
            blog.content = re.sub(pattern, img_md, blog.content, count=1)
        elif placeholder in blog.content:
            blog.content = blog.content.replace(placeholder, img_md)


    db.commit()
    
    return {
        "success": True,
        "url": f"/images/{image_filename}",
        "filename": image_filename,
        "updated_content": blog.content
    }


@app.post("/signup", response_model=Token)
def signup(user: UserCreate, db: Session = Depends(database.get_db)):
    db_user = db.query(models.User).filter(models.User.email == user.email).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    hashed_password = auth.get_password_hash(user.password)
    new_user = models.User(email=user.email, hashed_password=hashed_password)
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    
    access_token = auth.create_access_token(data={"sub": new_user.email})
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/login", response_model=Token)
def login(user: UserCreate, db: Session = Depends(database.get_db)):
    db_user = db.query(models.User).filter(models.User.email == user.email).first()
    if not db_user or not auth.verify_password(user.password, db_user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect email or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    access_token = auth.create_access_token(data={"sub": db_user.email})
    return {"access_token": access_token, "token_type": "bearer"}

@app.get("/me")
def get_me(current_user: models.User = Depends(get_current_user)):
    return {"email": current_user.email, "id": current_user.id}

@app.post("/forgot-password")
def forgot_password(request: ForgotPasswordRequest, db: Session = Depends(database.get_db)):
    user = db.query(models.User).filter(models.User.email == request.email).first()
    if not user:
        raise HTTPException(status_code=404, detail="Email not found")
    return {"message": "Email validated"}

@app.post("/reset-password")
def reset_password(request: PasswordResetConfirm, db: Session = Depends(database.get_db)):
    # In this simplified flow, the token field in PasswordResetConfirm will be used for email if we don't want to change the model, 
    # but it's better to update the PasswordResetConfirm model to use 'email'.
    # Actually, I'll update the Pydantic model first.
    user = db.query(models.User).filter(models.User.email == request.email).first()
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    user.hashed_password = auth.get_password_hash(request.new_password)
    db.commit()
    
    return {"message": "Password successfully reset"}

# --- Blog Writer Endpoints ---

def safe_slug(title: str) -> str:
    s = title.strip().lower()
    s = re.sub(r"[^a-z0-9 _-]+", "", s)
    s = re.sub(r"\s+", "_", s).strip("_")
    return s or "blog"

@app.get("/past-blogs")
def get_past_blogs(db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    db_blogs = db.query(models.Blog).filter(models.Blog.user_id == current_user.id).order_by(models.Blog.created_at.desc()).all()
    return [{"filename": b.filename, "title": b.title, "mtime": b.created_at.timestamp()} for b in db_blogs]

@app.get("/blog/{filename}")
def get_blog(filename: str, db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    blog = db.query(models.Blog).filter(models.Blog.filename == filename, models.Blog.user_id == current_user.id).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    return {"content": blog.content}

@app.get("/public/blog/{filename}")
def get_public_blog(filename: str, db: Session = Depends(database.get_db)):
    # Look up by filename regardless of user_id, since the filename should be unique (UUID-based).
    blog = db.query(models.Blog).filter(models.Blog.filename == filename).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    return {"content": blog.content, "title": blog.title}

class BlogUpdateRequest(BaseModel):
    content: str

@app.put("/blog/{filename}")
def update_blog(filename: str, request: BlogUpdateRequest, db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    blog = db.query(models.Blog).filter(models.Blog.filename == filename, models.Blog.user_id == current_user.id).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    
    blog.content = request.content
    blog.updated_at = datetime.utcnow()
    db.commit()
    return {"message": "Blog updated successfully"}

class BlogGenerateRequest(BaseModel):
    topic: str
    as_of: Optional[str] = None

@app.post("/generate-blog")
async def generate_blog(request: BlogGenerateRequest, db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    as_of = request.as_of or date.today().isoformat()
    
    inputs = {
        "topic": request.topic,
        "as_of": as_of,
        "recency_days": 7,
        "mode": "",
        "needs_research": False,
        "queries": [],
        "evidence": [],
        "plan": None,
        "sections": [],
        "merged_md": "",
        "md_with_placeholders": "",
        "image_specs": [],
        "generated_images": {},
        "final": "",
    }
    
    try:
        # Run the workflow
        result = blog_app.invoke(inputs)
        
        final_md = result.get("final", "")
        filename = None
        if final_md:
            from bwa_backend import _safe_slug
            title = result.get("topic")
            if not title and result.get("plan"):
                plan = result.get("plan")
                title = getattr(plan, "blog_title", None) if not isinstance(plan, dict) else plan.get("blog_title")
            
            filename = f"{_safe_slug(title or request.topic)}_{secrets.token_hex(4)}.md"
            
            # Save blog to database with a fresh session after long-running workflow
            session = database.SessionLocal()
            try:
                db_blog = session.query(models.Blog).filter(
                    models.Blog.filename == filename, 
                    models.Blog.user_id == current_user.id
                ).first()
                
                if db_blog:
                    db_blog.content = final_md
                    db_blog.title = title or request.topic
                else:
                    db_blog = models.Blog(
                        title=title or request.topic,
                        filename=filename,
                        content=final_md,
                        user_id=current_user.id
                    )
                    session.add(db_blog)
                
                session.commit()
                session.refresh(db_blog)

                # Save generated images to database
                generated_images = result.get("generated_images", {})
                for img_filename, img_data in generated_images.items():
                    db_image = session.query(models.BlogImage).filter(
                        models.BlogImage.filename == img_filename,
                        models.BlogImage.blog_id == db_blog.id
                    ).first()
                    if not db_image:
                        db_image = models.BlogImage(
                            filename=img_filename,
                            content=img_data,
                            blog_id=db_blog.id
                        )
                        session.add(db_image)
                    else:
                        db_image.content = img_data
                
                session.commit()
            except Exception as save_err:
                session.rollback()
                print(f" Failed to save blog to database: {save_err}")
                raise HTTPException(status_code=500, detail=f"Database error: Could not save blog ({str(save_err)})")
            finally:
                session.close()
            
        # IMPORTANT: Remove binary data before returning to frontend. 
        # FastAPI's jsonable_encoder tries to .decode() bytes to UTF-8, which crashes for image data.
        result.pop("generated_images", None)
        result["filename"] = filename
        return result
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error generating blog: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download-md/{filename}")
def download_md(filename: str, db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    blog = db.query(models.Blog).filter(models.Blog.filename == filename, models.Blog.user_id == current_user.id).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    
    return StreamingResponse(
        io.BytesIO(blog.content.encode("utf-8")),
        media_type="text/markdown",
        headers={"Content-Disposition": f"attachment; filename={blog.filename}"}
    )


def markdown_to_docx(md_text: str, blog_title: str, db: Session, blog_id: int) -> bytes:
    """
    Converts markdown to a .docx file with embedded images fetched from DB.
    Uses a simplified regex-based parser.
    """
    from docx import Document
    from docx.shared import Inches
    import io

    doc = Document()
    doc.add_heading(blog_title, 0)

    # Simplified parser for headings, images, and text
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            if line[2:].strip() != blog_title:
                doc.add_heading(line[2:].strip(), 1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), 2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), 3)
        
        # Images: ![alt](src)
        elif "![" in line and "](" in line:
            matches = list(re.finditer(r"!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)", line))
            if matches:
                last_end = 0
                for m in matches:
                    pre_text = line[last_end:m.start()].strip()
                    if pre_text:
                        p = doc.add_paragraph()
                        _add_formatted_runs(p, pre_text)
                    
                    src = m.group("src").strip().strip("<>").lstrip("./")
                    # Always just take the last part as the filename
                    img_filename = src.split("/")[-1]
                    
                    db_image = db.query(models.BlogImage).filter(
                        models.BlogImage.filename == img_filename,
                        models.BlogImage.blog_id == blog_id
                    ).first()
                    
                    caption = None
                    if i + 1 < len(lines):
                        next_line = lines[i+1].strip()
                        if next_line.startswith("*") and next_line.endswith("*"):
                            caption = next_line[1:-1].strip()
                            i += 1 # skip caption line
                    
                    if db_image:
                        try:
                            from PIL import Image
                            img_buffer = io.BytesIO(db_image.content)
                            out_buffer = io.BytesIO()
                            with Image.open(img_buffer) as img:
                                if img.mode in ("RGBA", "P"):
                                    img = img.convert("RGB")
                                img.save(out_buffer, format="JPEG", quality=85)
                            
                            out_buffer.seek(0)
                            doc.add_picture(out_buffer, width=Inches(5.5))
                            if caption:
                                doc.add_paragraph(caption, style='Caption')
                        except Exception as e:
                            error_type = type(e).__name__
                            print(f" DOCX Image Error ({error_type}): {e}")
                            doc.add_paragraph(f"[Image Error ({error_type}): {e}]")
                    else:
                        doc.add_paragraph(f"[Image not found: {src}]")
                    
                    last_end = m.end()
                
                post_text = line[last_end:].strip()
                if post_text:
                    p = doc.add_paragraph()
                    _add_formatted_runs(p, post_text)
            else:
                p = doc.add_paragraph()
                _add_formatted_runs(p, line)
        
        else:
            clean_line = re.sub(r"(?<!\!)\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", line)
            p = doc.add_paragraph()
            _add_formatted_runs(p, clean_line)
        
        i += 1

    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()

def _add_formatted_runs(paragraph, text: str):
    """Helper to add bold/italic runs to a python-docx paragraph."""
    # Split by ** for bold
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            # Split bold text by * for italic
            sub_parts = re.split(r'(\*.*?\*)', bold_text)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*'):
                    run = paragraph.add_run(sub_part[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    run = paragraph.add_run(sub_part)
                    run.bold = True
        else:
            # Split by * for italic
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*'):
                    run = paragraph.add_run(sub_part[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sub_part)

    return paragraph

    # We need to add the return statements back to markdown_to_docx


@app.get("/download-docx/{filename}")
def download_docx(filename: str, db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    blog = db.query(models.Blog).filter(models.Blog.filename == filename, models.Blog.user_id == current_user.id).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    
    docx_bytes = markdown_to_docx(blog.content, blog.title or blog.filename, db, blog.id)
    
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={safe_slug(blog.title or blog.filename)}.docx"}
    )


@app.get("/health")
def read_root():
    return {"status": "ok", "message": "Blog Writing Agent API is running"}


# ============== LinkedIn OAuth Flow ==============

LINKEDIN_CLIENT_ID = os.getenv("LINKEDIN_CLIENT_ID")
LINKEDIN_CLIENT_SECRET = os.getenv("LINKEDIN_CLIENT_SECRET")
LINKEDIN_REDIRECT_URI = os.getenv("LINKEDIN_REDIRECT_URI", "http://localhost:8000/linkedin/callback")
LINKEDIN_SCOPES = "openid profile w_member_social"


@app.get("/linkedin/auth")
def linkedin_auth(token: str = None):
    """Step 1: Redirect user to LinkedIn's authorization page."""
    if not LINKEDIN_CLIENT_ID:
        raise HTTPException(status_code=500, detail="LinkedIn Client ID not configured")
    
    # Store the app's JWT token in state so we can identify the user after callback
    import urllib.parse
    state = token or ""
    auth_url = (
        f"https://www.linkedin.com/oauth/v2/authorization?"
        f"response_type=code&"
        f"client_id={LINKEDIN_CLIENT_ID}&"
        f"redirect_uri={urllib.parse.quote(LINKEDIN_REDIRECT_URI)}&"
        f"scope={urllib.parse.quote(LINKEDIN_SCOPES)}&"
        f"state={state}"
    )
    from fastapi.responses import RedirectResponse
    return RedirectResponse(url=auth_url)


@app.get("/linkedin/callback")
def linkedin_callback(code: str = None, state: str = None, error: str = None, db: Session = Depends(database.get_db)):
    """Step 2: LinkedIn redirects back here with an auth code. Exchange it for an access token."""
    if error:
        return {"error": error}
    if not code:
        raise HTTPException(status_code=400, detail="No authorization code received")
    
    # Exchange code for access token
    token_resp = requests.post("https://www.linkedin.com/oauth/v2/accessToken", data={
        "grant_type": "authorization_code",
        "code": code,
        "redirect_uri": LINKEDIN_REDIRECT_URI,
        "client_id": LINKEDIN_CLIENT_ID,
        "client_secret": LINKEDIN_CLIENT_SECRET,
    }, headers={"Content-Type": "application/x-www-form-urlencoded"})
    
    if not token_resp.ok:
        raise HTTPException(status_code=400, detail=f"Failed to get token: {token_resp.text}")
    
    linkedin_token = token_resp.json().get("access_token")
    
    # Get user info from LinkedIn
    userinfo_resp = requests.get("https://api.linkedin.com/v2/userinfo", headers={"Authorization": f"Bearer {linkedin_token}"})
    if not userinfo_resp.ok:
        raise HTTPException(status_code=400, detail=f"Failed to get user info: {userinfo_resp.text}")
    
    person_urn = f"urn:li:person:{userinfo_resp.json()['sub']}"
    linkedin_name = userinfo_resp.json().get("name", "Unknown")
    
    # Find user by JWT token in state
    if state:
        try:
            payload = auth.jwt.decode(state, auth.SECRET_KEY, algorithms=[auth.ALGORITHM])
            email = payload.get("sub")
            user = db.query(models.User).filter(models.User.email == email).first()
            if user:
                user.linkedin_token = linkedin_token
                user.linkedin_person_urn = person_urn
                db.commit()
        except Exception as e:
            print(f"Failed to decode state token: {e}")
    
    # Return a nice HTML page that closes the popup
    html = f"""
    <html>
    <body style="font-family: Arial; text-align: center; padding-top: 100px; background: #1a1a2e; color: white;">
        <h2> LinkedIn Connected!</h2>
        <p>Connected as: <strong>{linkedin_name}</strong></p>
        <p>You can close this window now.</p>
        <script>
            if (window.opener) {{
                window.opener.postMessage({{ type: 'LINKEDIN_CONNECTED', name: '{linkedin_name}' }}, '*');
                setTimeout(() => window.close(), 2000);
            }}
        </script>
    </body>
    </html>
    """
    from fastapi.responses import HTMLResponse
    return HTMLResponse(content=html)


@app.get("/linkedin/status")
def linkedin_status(db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    """Check if the current user has connected their LinkedIn account."""
    if current_user.linkedin_token:
        # Verify the token is still valid
        resp = requests.get("https://api.linkedin.com/v2/userinfo", headers={"Authorization": f"Bearer {current_user.linkedin_token}"})
        if resp.ok:
            name = resp.json().get("name", "Connected")
            return {"connected": True, "name": name}
        else:
            # Token expired, clear it
            current_user.linkedin_token = None
            current_user.linkedin_person_urn = None
            db.commit()
    return {"connected": False}


@app.post("/linkedin/disconnect")
def linkedin_disconnect(db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    """Disconnect LinkedIn from the current user's account."""
    current_user.linkedin_token = None
    current_user.linkedin_person_urn = None
    db.commit()
    return {"message": "LinkedIn disconnected"}


class LinkedInPostRequest(BaseModel):
    access_token: Optional[str] = None

@app.post("/blog/{filename}/post-to-linkedin")
def post_to_linkedin(filename: str, req: LinkedInPostRequest = LinkedInPostRequest(), db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    blog = db.query(models.Blog).filter(models.Blog.filename == filename, models.Blog.user_id == current_user.id).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")

    # Use OAuth token if connected, otherwise fallback to manually provided token
    token = current_user.linkedin_token or (req.access_token if req else None)
    if not token:
        raise HTTPException(status_code=400, detail="LinkedIn not connected. Please connect your LinkedIn account first.")

    person_urn = current_user.linkedin_person_urn
    headers = {
        'Authorization': f'Bearer {token}',
        'X-Restli-Protocol-Version': '2.0.0',
        'Content-Type': 'application/json'
    }

    # If we don't have the person URN stored, fetch it
    if not person_urn:
        me_resp = requests.get('https://api.linkedin.com/v2/userinfo', headers=headers)
        if not me_resp.ok:
            raise HTTPException(status_code=400, detail=f"LinkedIn Auth failed: {me_resp.text}")
        person_urn = f"urn:li:person:{me_resp.json()['sub']}"
        current_user.linkedin_person_urn = person_urn
        db.commit()

    # --- Build Eye-Catching LinkedIn Article Post ---
    blog_title = blog.title or request.topic if 'request' in locals() else (blog.title or "Blog Post")
    raw_blog_content = blog.content or ""
    
    # Get the public base URL (Render domain or tunnel)
    public_base_url = os.getenv("PUBLIC_URL", "https://blog-creation-eetn.onrender.com")
    blog_url = f"{public_base_url}/p/{filename}"
    
    # Extract clean text for summary
    clean_text = re.sub(r'!\[.*?\]\([^)]+\)', '', raw_blog_content)
    clean_text = re.sub(r'^\s*\*[^*]+\*\s*$', '', clean_text, flags=re.MULTILINE)
    clean_text = re.sub(r'\[\[IMAGE_\d+\]\]', '', clean_text)
    clean_text = re.sub(r'^#+ .*$', '', clean_text, flags=re.MULTILINE)
    clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', clean_text)
    clean_text = re.sub(r'\*(.+?)\*', r'\1', clean_text)
    clean_text = re.sub(r'\n{2,}', '\n\n', clean_text).strip()
    
    # Extract first 2-3 sentences for a strong hook
    sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', clean_text) if len(s.strip()) > 10]
    hook = " ".join(sentences[:2]) if len(sentences) >= 2 else (sentences[0] if sentences else clean_text[:200])
    if len(hook) > 250:
        hook = hook[:245] + "..."
        
    # Generate auto-hashtags based on title words
    title_words = [re.sub(r'[^a-zA-Z0-9]', '', w) for w in blog_title.split()]
    title_words = [w.capitalize() for w in title_words if len(w) > 3]
    hashtags = " ".join([f"#{w}" for w in title_words[:5]])
    if not hashtags:
        hashtags = "#Tech #Innovation #AI #Article"
    
    # Build professional, eye-catching post commentary
    commentary = (
        f"🚀 {blog_title}\n\n"
        f"{hook}\n\n"
        f"Discover full insights, key architecture, and practical takeaways in the complete article below!\n\n"
        f"{hashtags}"
    )
    
    # Enforce LinkedIn 3,000 character limit
    if len(commentary) > 3000:
        commentary = commentary[:2990] + '...'

    # Add a unique query param to prevent LinkedIn duplicate post rejection
    unique_blog_url = f"{blog_url}?ref={secrets.token_hex(3)}"

    # --- Post as a Rich ARTICLE Share Card (Like LinkedIn Newsletters & Top Companies) ---
    media_item = {
        "status": "READY",
        "originalUrl": unique_blog_url,
        "title": {"text": blog_title},
        "description": {"text": hook[:200]}
    }
    
    # Attach explicit thumbnail image URL if the blog has an image
    first_image = blog.images[0] if (hasattr(blog, 'images') and blog.images) else None
    if first_image:
        img_url = f"{public_base_url.rstrip('/')}/images/{first_image.filename}"
        media_item["thumbnails"] = [{"url": img_url}]

    post_data = {
        "author": person_urn,
        "lifecycleState": "PUBLISHED",
        "specificContent": {
            "com.linkedin.ugc.ShareContent": {
                "shareCommentary": {"text": commentary},
                "shareMediaCategory": "ARTICLE",
                "media": [media_item]
            }
        },
        "visibility": {"com.linkedin.ugc.MemberNetworkVisibility": "PUBLIC"}
    }
        
    post_resp = requests.post('https://api.linkedin.com/v2/ugcPosts', headers=headers, json=post_data)
    
    if not post_resp.ok:
        print("LinkedIn API error response:", post_resp.text)
        raise HTTPException(status_code=400, detail=f"Failed to publish post: {post_resp.text}")
        
    post_json = post_resp.json() if post_resp.content else {}
    post_urn = post_json.get("id", "")
    
    has_image = bool(first_image)
    post_type = "article link with image" if has_image else "article link"
    return {
        "message": f"Successfully posted to LinkedIn as {post_type}!",
        "post_id": post_urn,
        "blog_url": blog_url,
        "has_image": has_image,
        "url": "https://www.linkedin.com/in/me/recent-activity/all/"
    }

# --- Static File Serving (For Production Monolith) ---
@app.get("/{file_name:path}")
def serve_react_app(file_name: str):
    dist_path = Path("frontend/dist")
    file_path = dist_path / file_name
    
    if file_path.is_file():
        return FileResponse(file_path)
    
    index_path = dist_path / "index.html"
    if index_path.is_file():
        return FileResponse(index_path)
        
    from fastapi.responses import HTMLResponse
    return HTMLResponse("<h1>Frontend not built yet. Run 'npm run build' in the frontend directory.</h1>", status_code=404)
